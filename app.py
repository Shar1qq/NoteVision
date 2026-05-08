"""
NoteVision – Intelligent Handwriting Converter
===============================================
AI-powered system that perceives, assesses, converts, critiques and refines
handwritten notes into structured digital documents.

UI layers:
  - Agent Activity Log   : real-time step-by-step transparency
  - Assessment Panel     : what the agent perceived
  - Strategy Badge       : which tool the agent chose and why
  - Critique Panel       : self-review output
  - Confidence Meter     : agent certainty score
  - Human-in-the-Loop    : user can accept, reject, or request another pass
  - Memory Dashboard     : session + history stats
"""

import streamlit as st
from google import genai
from google.oauth2 import service_account
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import re
import streamlit.components.v1 as components

from dataclasses import asdict
from agent import AgentMemory, NoteVisionAgent, SafetyResult
from PIL import Image

# ─────────────────────────────────────────────────────────────────────────────
# Page Config
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="NoteVision",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
# Styling
# ─────────────────────────────────────────────────────────────────────────────

st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

  html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

  .stApp { background: #0f1117; }

  /* ── Header ── */
  .nv-header {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    border-radius: 16px;
    padding: 2rem 2.5rem;
    margin-bottom: 1.5rem;
  }
  .nv-header h1 {
    color: white !important;
    font-size: 2rem !important;
    font-weight: 800 !important;
    margin: 0 !important;
    -webkit-text-fill-color: white !important;
  }
  .nv-header p { color: rgba(255,255,255,0.85); margin: 0.4rem 0 0 0; font-size: 1rem; }

  /* ── Cards ── */
  .nv-card {
    background: #1e2130;
    border: 1px solid #2d3250;
    border-radius: 12px;
    padding: 1.25rem 1.5rem;
    margin-bottom: 1rem;
  }
  .nv-card-title {
    font-size: 0.75rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: #667eea;
    margin-bottom: 0.6rem;
  }

  /* ── Agent log ── */
  .agent-log {
    background: #0d1117;
    border: 1px solid #21262d;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    font-family: 'Courier New', monospace;
    font-size: 0.82rem;
    color: #8b949e;
    max-height: 220px;
    overflow-y: auto;
    line-height: 1.7;
  }
  .agent-log .step { color: #58a6ff; }
  .agent-log .ok   { color: #3fb950; }
  .agent-log .warn { color: #d29922; }

  /* ── Strategy badge ── */
  .strategy-badge {
    display: inline-block;
    background: linear-gradient(135deg, #667eea22, #764ba222);
    border: 1px solid #667eea55;
    color: #a5b4fc;
    border-radius: 20px;
    padding: 0.3rem 1rem;
    font-size: 0.85rem;
    font-weight: 600;
  }

  /* ── Confidence meter ── */
  .conf-bar-wrap {
    background: #0d1117;
    border-radius: 8px;
    height: 10px;
    overflow: hidden;
    margin-top: 0.5rem;
  }
  .conf-bar-fill {
    height: 100%;
    border-radius: 8px;
    background: linear-gradient(90deg, #667eea, #764ba2);
    transition: width 0.4s ease;
  }

  /* ── Buttons ── */
  .stButton > button {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white;
    border: none;
    border-radius: 10px;
    padding: 0.6rem 1.5rem;
    font-weight: 600;
    font-size: 0.95rem;
    transition: all 0.2s ease;
    box-shadow: 0 4px 14px rgba(102, 126, 234, 0.35);
  }
  .stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 20px rgba(102, 126, 234, 0.55);
  }
  .stDownloadButton > button {
    background: #1e2130;
    color: #667eea;
    border: 1.5px solid #667eea;
    border-radius: 10px;
    font-weight: 600;
    transition: all 0.2s;
  }
  .stDownloadButton > button:hover {
    background: #667eea;
    color: white;
  }

  /* ── Sidebar ── */
  [data-testid="stSidebar"] { background: #161b27; border-right: 1px solid #2d3250; }

  /* ── Tabs ── */
  .stTabs [data-baseweb="tab-list"] {
    background: #1e2130;
    border-radius: 10px;
    padding: 4px;
    gap: 4px;
  }
  .stTabs [data-baseweb="tab"] {
    background: transparent;
    border-radius: 8px;
    color: #8892b0;
    font-weight: 500;
    font-size: 0.85rem;
  }
  .stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #667eea, #764ba2) !important;
    color: white !important;
  }

  /* ── Assessment grid ── */
  .assess-grid {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 0.75rem;
    margin-top: 0.5rem;
  }
  .assess-item {
    background: #0d1117;
    border: 1px solid #21262d;
    border-radius: 8px;
    padding: 0.6rem 0.9rem;
    text-align: center;
  }
  .assess-label { font-size: 0.7rem; color: #6e7b8b; text-transform: uppercase; letter-spacing: 0.06em; }
  .assess-value { font-size: 1rem; font-weight: 700; color: #e2e8f0; margin-top: 2px; }

  /* ── History item ── */
  .hist-item {
    background: #0d1117;
    border: 1px solid #21262d;
    border-radius: 8px;
    padding: 0.6rem 1rem;
    margin-bottom: 0.4rem;
    font-size: 0.82rem;
    color: #8b949e;
  }

  /* ── HITL controls ── */
  .hitl-banner {
    background: linear-gradient(135deg, #1a2a1a, #1e2a1e);
    border: 1.5px solid #2ea04388;
    border-radius: 10px;
    padding: 1rem 1.3rem;
    margin: 0.75rem 0;
  }
  .hitl-title { color: #3fb950; font-weight: 700; font-size: 0.9rem; margin-bottom: 0.3rem; }
  .hitl-desc  { color: #8b949e; font-size: 0.82rem; }

  hr { border: none; border-top: 1px solid #2d3250; margin: 1.2rem 0; }

  /* ── Pipeline visualization ── */
  .pipeline-container { margin: 0.25rem 0; }
  .pipeline-step {
    display: flex;
    align-items: center;
    padding: 0.65rem 1rem;
    margin-bottom: 0.4rem;
    border-radius: 10px;
    border: 1px solid #21262d;
    background: #0d1117;
    transition: all 0.3s ease;
  }
  .ps-done  { border-color: #2ea04344 !important; background: #0a160a !important; }
  .ps-active { border-color: #667eea88 !important; background: linear-gradient(135deg,#667eea0d,#764ba20d) !important; }
  .ps-pending { opacity: 0.3; }
  .ps-icon { font-size: 1.15rem; width: 2rem; text-align: center; flex-shrink: 0; }
  .ps-content { flex: 1; margin-left: 0.75rem; }
  .ps-name { font-size: 0.83rem; font-weight: 700; color: #c9d1d9; }
  .ps-msg  { font-size: 0.73rem; color: #6e7b8b; font-family: 'Courier New', monospace; margin-top: 2px; }
  .ps-status { margin-left: 1rem; flex-shrink: 0; display: flex; align-items: center; }
  @keyframes nv-pulse {
    0%,100% { opacity:1; transform:scale(1); }
    50%     { opacity:0.3; transform:scale(0.7); }
  }
  .ps-spinner {
    width: 10px; height: 10px;
    background: #667eea;
    border-radius: 50%;
    display: inline-block;
    animation: nv-pulse 1.1s ease-in-out infinite;
  }
  .ps-connector {
    width: 2px; height: 10px;
    background: #21262d;
    margin: 0 auto 0 1.9rem;
  }

  /* ── Safety gate ── */
  .safety-gate {
    background: linear-gradient(135deg, #2a1a0a, #2a200a);
    border: 1.5px solid #d2992288;
    border-radius: 12px;
    padding: 1.25rem 1.5rem;
    margin: 1rem 0;
  }
  .safety-title { color: #d29922; font-weight: 700; font-size: 1rem; margin-bottom: 0.6rem; }
  .safety-desc  { color: #c9d1d9; font-size: 0.86rem; line-height: 1.75; }

  /* ── Consent banner ── */
  .consent-banner {
    background: linear-gradient(135deg, #0d1117, #161b27);
    border: 1px solid #2d3250;
    border-radius: 12px;
    padding: 1.5rem 2rem;
    margin: 1rem 0;
  }
  .consent-title { color: #c9d1d9; font-weight: 700; font-size: 1rem; margin-bottom: 0.75rem; }
  .consent-item  { font-size: 0.83rem; color: #8b949e; margin-bottom: 0.4rem; }

  /* ── Output mode toggle ── */
  .mode-toggle-card {
    background: linear-gradient(135deg, #0d1117, #161b27);
    border: 1px solid #2d3250;
    border-radius: 14px;
    padding: 0.9rem 1.3rem;
    margin: 0.8rem 0 1rem 0;
    display: flex;
    align-items: center;
    gap: 1rem;
  }
  .mode-label {
    font-size: 0.78rem;
    font-weight: 600;
    color: #c9d1d9;
    line-height: 1.3;
  }
  .mode-desc { font-size: 0.71rem; color: #6e7b8b; margin-top: 0.15rem; }

  /* ── IPR notice ── */
  .ipr-notice {
    background: #0d1117;
    border-left: 3px solid #667eea55;
    border-radius: 0 6px 6px 0;
    padding: 0.5rem 0.9rem;
    margin-bottom: 0.75rem;
    font-size: 0.76rem;
    color: #4a5568;
  }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Rate Limiting
# ─────────────────────────────────────────────────────────────────────────────

MAX_PER_SESSION = 20
COOLDOWN = 10

if "conversion_count" not in st.session_state:
    st.session_state.conversion_count = 0
if "last_time" not in st.session_state:
    st.session_state.last_time = 0
if "agent_result" not in st.session_state:
    st.session_state.agent_result = None
if "hitl_accepted" not in st.session_state:
    st.session_state.hitl_accepted = None   # None | True | False
if "diagram_mode" not in st.session_state:
    st.session_state.diagram_mode = True
if "consent_given" not in st.session_state:
    st.session_state.consent_given = False
if "safety_result" not in st.session_state:
    st.session_state.safety_result = None
if "safety_ack" not in st.session_state:
    st.session_state.safety_ack = False
if "pending_run" not in st.session_state:
    st.session_state.pending_run = False

# ─────────────────────────────────────────────────────────────────────────────
# API + Agent init
# ─────────────────────────────────────────────────────────────────────────────

try:
    credentials = service_account.Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=["https://www.googleapis.com/auth/cloud-platform"]
    )
    gemini_client = genai.Client(
        vertexai=True,
        project=st.secrets["gcp_service_account"]["project_id"],
        location="us-central1",
        credentials=credentials
    )
except Exception as e:
    st.error(f"⚠️ Vertex AI credentials not configured. Error: {e}")
    st.stop()

memory = AgentMemory(st.session_state)
agent = NoteVisionAgent(gemini_client, memory)

# ─────────────────────────────────────────────────────────────────────────────
# Header
# ─────────────────────────────────────────────────────────────────────────────

# ── Privacy Consent Gate (shown once per session) ───────────────────────────
if not st.session_state.consent_given:
    st.markdown("""
<div class="consent-banner">
  <div class="consent-title">🔒 Privacy & Responsible Use Notice</div>
  <div class="consent-item">🔵 <b style="color:#c9d1d9;">Data Processing</b> — Images are sent to Google Gemini 2.5 Flash via Vertex AI (GCP). No images or personal data are stored permanently on any server.</div>
  <div class="consent-item">🔵 <b style="color:#c9d1d9;">Intellectual Property</b> — Only upload content you own or have explicit rights to process. Uploading third-party copyrighted material without authorisation may violate IPR law.</div>
  <div class="consent-item">🔵 <b style="color:#c9d1d9;">AI Limitations</b> — This system is semi-autonomous. All output requires human review and approval before use.</div>
  <div class="consent-item">🔵 <b style="color:#c9d1d9;">Ethical Use</b> — Do not use this system to process confidential, sensitive, or exam-related material without appropriate authorisation.</div>
  <div style="margin-top:1rem;font-size:0.78rem;color:#4a5568;">By continuing, you acknowledge these terms and agree to responsible use of this AI system.</div>
</div>
""", unsafe_allow_html=True)
    if st.button("✔️ I understand and agree — Continue to NoteVision", type="primary"):
        st.session_state.consent_given = True
        st.rerun()
    st.stop()

st.markdown("""
<div class="nv-header">
  <h1>📝 NoteVision</h1>
  <p>Intelligent Handwriting Converter &nbsp;·&nbsp; Powered by Google Gemini</p>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Sidebar – Memory Dashboard
# ─────────────────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("### 🧠 Memory Dashboard")
    st.markdown(f"""
<div class="nv-card">
  <div class="nv-card-title">Session Memory</div>
  <div style="font-size:2rem;font-weight:800;color:#667eea;">
    {memory.session_conversions()}
    <span style="font-size:1rem;color:#6e7b8b;">/ {MAX_PER_SESSION}</span>
  </div>
  <div style="font-size:0.8rem;color:#6e7b8b;margin-top:4px;">conversions this session</div>
</div>
<div class="nv-card">
  <div class="nv-card-title">Long-term Memory</div>
  <div style="font-size:1.6rem;font-weight:700;color:#a5b4fc;">
    {memory.total_conversions()}
  </div>
  <div style="font-size:0.8rem;color:#6e7b8b;margin-top:4px;">total conversions ever</div>
</div>
""", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ⚙️ Agent Architecture")
    st.markdown("""
<div style="font-size:0.82rem;color:#8b949e;line-height:1.9;">
� <b style="color:#d29922;">[Step 0]</b> <b style="color:#c9d1d9;">Safety Check</b> – ethical screening<br>
� <b style="color:#c9d1d9;">Perceive</b> – quality assessment tool<br>
🔵 <b style="color:#c9d1d9;">Decide</b> &nbsp;– strategy selector (rule-based)<br>
🔵 <b style="color:#c9d1d9;">Act</b> &nbsp;&nbsp;&nbsp;&nbsp;– Gemini Vision conversion<br>
🔵 <b style="color:#c9d1d9;">Critique</b> – self-review tool<br>
🔵 <b style="color:#c9d1d9;">Refine</b> &nbsp;– feedback-based correction<br>
🟢 <b style="color:#3fb950;">HITL ×2</b> &nbsp;– safety gate + output approval
</div>
""", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ⚖️ Ethics & Compliance")
    st.markdown("""
<div style="font-size:0.79rem;color:#8b949e;line-height:1.85;">
<b style="color:#667eea;">ACM/IEEE Principles Applied:</b><br>
🔵 <b style="color:#c9d1d9;">§1.1 Public Good</b> — accessibility<br>
🔵 <b style="color:#c9d1d9;">§1.6 Privacy</b> — no data stored<br>
🔵 <b style="color:#c9d1d9;">§2.5 Evaluation</b> — critique loop<br>
🔵 <b style="color:#c9d1d9;">§2.9 Inclusion</b> — safety gate<br>
🔵 <b style="color:#c9d1d9;">§3.1 Limitations</b> — confidence score<br>
<span style="color:#d29922;">⚠️ §2.5 Bias risk</span> — mitigated via confidence &amp; HITL
</div>
""", unsafe_allow_html=True)

    st.markdown("---")
    # Recent history
    history = memory.get_history()[-5:]
    if history:
        st.markdown("### 📂 Recent History")
        for h in reversed(history):
            q = h["assessment"]["quality"]
            colour = {"good": "#3fb950", "poor": "#f85149", "unclear": "#d29922"}.get(q, "#8b949e")
            st.markdown(f"""
<div class="hist-item">
  <b style="color:#c9d1d9;">{h['filename'][:22]}</b><br>
  <span style="color:{colour};">●</span> {q} &nbsp;·&nbsp; {h['strategy_used']} &nbsp;·&nbsp;
  {h['confidence']:.0%} conf<br>
  <span style="color:#4a5568;">{h['timestamp']}</span>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Main – File Upload
# ─────────────────────────────────────────────────────────────────────────────

uploaded = st.file_uploader(
    "Upload handwritten notes (JPG / PNG / JPEG)",
    type=["jpg", "jpeg", "png"],
    help="The agent will automatically assess quality and choose the best strategy.",
)

if uploaded is not None:
    image = Image.open(uploaded)

    col_img, col_info = st.columns([1, 1])
    with col_img:
        with st.expander("📸 Original Image", expanded=True):
            st.image(image, use_container_width=True)

    with col_info:
        st.markdown("""
<div class="nv-card">
  <div class="nv-card-title">How the Agent Works</div>
  <div style="font-size:0.85rem;color:#8b949e;line-height:1.9;">
    1️⃣ <b style="color:#c9d1d9;">Perceive</b> — analyses image quality & content<br>
    2️⃣ <b style="color:#c9d1d9;">Decide</b> — picks the best conversion strategy<br>
    3️⃣ <b style="color:#c9d1d9;">Act</b> — converts with strategy-specific prompt<br>
    4️⃣ <b style="color:#c9d1d9;">Critique</b> — self-reviews its own output<br>
    5️⃣ <b style="color:#c9d1d9;">Refine</b> — fixes issues found in critique<br>
    6️⃣ <b style="color:#c9d1d9;">HITL gate</b> — you approve before download
  </div>
</div>
""", unsafe_allow_html=True)

    # ─ Rate-limit check ───────────────────────────────────────────────────────
    import time as _time
    if st.session_state.conversion_count >= MAX_PER_SESSION:
        st.error(f"Session limit reached ({MAX_PER_SESSION} conversions). Refresh to start a new session.")
        st.stop()

    st.markdown('<div class="ipr-notice">⚖️ <b>IPR Notice:</b> Only upload content you own or have rights to process. Processing third-party copyrighted material without authorisation may violate intellectual property law.</div>', unsafe_allow_html=True)

    # ── Output Mode Toggle ───────────────────────────────────────────────────────
    col_mode_toggle, col_mode_info = st.columns([1, 5])
    with col_mode_toggle:
        new_mode = st.toggle(
            "Diagrams",
            value=st.session_state.diagram_mode,
            key="pre_run_mode_toggle",
            help="ON = render graphs as interactive Mermaid charts · OFF = LLM explains graphs in prose"
        )
        st.session_state.diagram_mode = new_mode
    with col_mode_info:
        if st.session_state.diagram_mode:
            st.markdown(
                "<div class='mode-label'>📊 Diagram Mode</div>"
                "<div class='mode-desc'>Graphs will be rendered as interactive Mermaid charts (xychart-beta)</div>",
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                "<div class='mode-label'>📝 Text Explanation Mode</div>"
                "<div class='mode-desc'>LLM will describe each graph in detailed prose — no charts generated</div>",
                unsafe_allow_html=True
            )

    # ── Ethical Safety Gate ─────────────────────────────────────────────────────
    if (st.session_state.safety_result and
            not st.session_state.safety_result.get("is_safe") and
            not st.session_state.safety_ack):
        sr = st.session_state.safety_result
        concern_labels = {
            "exam_paper":    "Academic Integrity — Exam / Test Paper Detected",
            "personal_data": "Data Privacy — Personal Records Detected",
            "confidential":  "Confidentiality — Restricted Document Detected",
            "inappropriate": "Content Policy — Inappropriate Material Detected",
        }
        concern_title = concern_labels.get(sr.get("concern_type", ""), "Sensitive Content Detected")
        st.markdown(f"""
<div class="safety-gate">
  <div class="safety-title">⚠️ Ethical Safety Notice — Human Review Required</div>
  <div class="safety-desc">
    <b style="color:#d29922;">{concern_title}</b><br><br>
    {sr.get("explanation", "")}<br><br>
    <b>Recommendation:</b> {sr.get("recommendation", "")}
  </div>
</div>""", unsafe_allow_html=True)
        col_ack, col_cancel = st.columns(2)
        with col_ack:
            if st.button("✋ I Acknowledge & Take Responsibility", type="primary", use_container_width=True):
                st.session_state.safety_ack = True
                st.rerun()
        with col_cancel:
            if st.button("🚫 Cancel — Do Not Process", use_container_width=True):
                st.session_state.safety_result = None
                st.session_state.safety_ack = False
                st.session_state.pending_run = False
                st.rerun()

    col_btn, _ = st.columns([1, 3])
    with col_btn:
        run_agent = st.button("🚀 Run Agent", type="primary", use_container_width=True)

    safety_cleared = (
        st.session_state.safety_result is not None and
        (st.session_state.safety_result.get("is_safe") or st.session_state.safety_ack)
    )
    should_execute = st.session_state.pending_run and safety_cleared

    if run_agent or should_execute:
        elapsed = _time.time() - st.session_state.last_time
        if elapsed < COOLDOWN:
            st.warning(f"⏳ Cooldown: please wait {int(COOLDOWN - elapsed)}s before the next run.")
        else:
            if run_agent:
                st.session_state.hitl_accepted = None
                st.session_state.agent_result = None
                st.session_state.safety_result = None
                st.session_state.safety_ack = False
                st.session_state.pending_run = True
                with st.spinner("🛡️ Ethical safety screening…"):
                    safety = agent.tools.safety_check(image)
                st.session_state.safety_result = asdict(safety)
                if not safety.is_safe:
                    st.rerun()

            st.session_state.pending_run = False
            log_lines = []

            PIPELINE_STEPS_UI = [
                ("�️", "Safety Check", "Ethical content screening"),
                ("��", "Perceive",     "Image quality & content analysis"),
                ("🧠", "Decide",       "Strategy selection"),
                ("⚙️", "Act",          "Gemini Vision conversion"),
                ("🔎", "Critique",     "Self-review pass"),
                ("✨", "Refine",       "Feedback-based correction"),
            ]

            def _get_pipeline_state(lines):
                txt = "\n".join(lines)
                if "Done. Confidence" in txt:        return list(range(6)), None
                elif "Applying refinements" in txt:  return [0,1,2,3,4], 5
                elif "💬 [Agent] Critique" in txt:   return [0,1,2,3,4], 5
                elif "Running self-critique" in txt: return [0,1,2,3], 4
                elif "Converting image" in txt:      return [0,1,2], 3
                elif "Strategy selected" in txt:     return [0,1,2], 3
                elif "Deciding conversion" in txt:   return [0,1], 2
                elif "Assessment complete" in txt:   return [0,1], 2
                elif "Perceiving image" in txt:      return [0], 1
                return [0], None

            def _build_pipeline_html(lines):
                done, active = _get_pipeline_state(lines)
                last_msg = lines[-1] if lines else ""
                rows = []
                for i, (icon, name, desc) in enumerate(PIPELINE_STEPS_UI):
                    if i in done:
                        cls = "ps-done"
                        status_html = '<span style="color:#3fb950;font-size:1.1rem;">✓</span>'
                        msg = desc
                    elif i == active:
                        cls = "ps-active"
                        status_html = '<span class="ps-spinner"></span>'
                        clean = last_msg.split("] ", 1)[-1] if "] " in last_msg else last_msg
                        msg = (clean[:88] + "…") if len(clean) > 88 else clean
                    else:
                        cls = "ps-pending"
                        status_html = '<span style="color:#3d4559;font-size:0.9rem;">○</span>'
                        msg = desc
                    connector = '<div class="ps-connector"></div>' if i < len(PIPELINE_STEPS_UI) - 1 else ""
                    rows.append(
                        f'<div class="pipeline-step {cls}">'
                        f'<div class="ps-icon">{icon}</div>'
                        f'<div class="ps-content"><div class="ps-name">{name}</div>'
                        f'<div class="ps-msg">{msg}</div></div>'
                        f'<div class="ps-status">{status_html}</div></div>{connector}'
                    )
                return '<div class="pipeline-container">' + "".join(rows) + "</div>"

            st.markdown('<div class="nv-card-title" style="margin-bottom:0.5rem;">⚡ Processing Pipeline</div>', unsafe_allow_html=True)
            log_placeholder = st.empty()

            def _update_log(msg: str):
                log_lines.append(msg)
                log_placeholder.markdown(_build_pipeline_html(log_lines), unsafe_allow_html=True)

            try:
                with st.spinner("⚙️ Processing…"):
                    result = agent.run(
                        image,
                        filename=uploaded.name,
                        log_callback=_update_log,
                        safety_result=st.session_state.safety_result,
                        diagram_mode=st.session_state.diagram_mode
                    )
                st.session_state.agent_result = result
                st.session_state.conversion_count += 1
                st.session_state.last_time = _time.time()
                st.rerun()
            except Exception as e:
                st.error(f"Processing failed: {e}")

# ─────────────────────────────────────────────────────────────────────────────
# Results Panel (after agent run)
# ─────────────────────────────────────────────────────────────────────────────

result = st.session_state.get("agent_result")
if result is not None:

    st.markdown("---")
    st.markdown("## 📊 Conversion Report")

    # ── Row 1: Assessment + Strategy + Confidence ─────────────────────────────
    c1, c2, c3 = st.columns([2, 1.5, 1.5])

    with c1:
        q = result.assessment.quality
        q_colour = {"good": "#3fb950", "poor": "#f85149", "unclear": "#d29922"}.get(q, "#8b949e")
        st.markdown(f"""
<div class="nv-card">
  <div class="nv-card-title">🔍 Perception — Image Assessment</div>
  <div class="assess-grid">
    <div class="assess-item">
      <div class="assess-label">Quality</div>
      <div class="assess-value" style="color:{q_colour};">{q.upper()}</div>
    </div>
    <div class="assess-item">
      <div class="assess-label">Has Math</div>
      <div class="assess-value">{"✓" if result.assessment.has_math else "✗"}</div>
    </div>
    <div class="assess-item">
      <div class="assess-label">Diagrams</div>
      <div class="assess-value">{"✓" if result.assessment.has_diagrams else "✗"}</div>
    </div>
    <div class="assess-item">
      <div class="assess-label">Density</div>
      <div class="assess-value" style="font-size:0.85rem;">{result.assessment.handwriting_density}</div>
    </div>
    <div class="assess-item" style="grid-column:span 2;">
      <div class="assess-label">Agent Notes</div>
      <div style="font-size:0.78rem;color:#8b949e;margin-top:4px;">{result.assessment.notes}</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

    with c2:
        strategy_labels = {
            "math_focused":         "📐 Math Focused",
            "careful_reconstruction": "🔬 Careful Reconstruction",
            "structure_aware":      "🗂 Structure Aware",
            "detail_preserving":    "🔎 Detail Preserving",
            "standard":             "⚡ Standard",
        }
        badge = strategy_labels.get(result.strategy_used, result.strategy_used)
        st.markdown(f"""
<div class="nv-card" style="height:100%;">
  <div class="nv-card-title">🧠 Decision — Strategy</div>
  <div style="margin-top:0.8rem;">
    <div class="strategy-badge">{badge}</div>
  </div>
  <div style="font-size:0.78rem;color:#6e7b8b;margin-top:0.8rem;">
    Selected automatically based on image perception.
  </div>
</div>
""", unsafe_allow_html=True)

    with c3:
        conf_pct = int(result.confidence * 100)
        conf_colour = "#3fb950" if conf_pct >= 75 else "#d29922" if conf_pct >= 50 else "#f85149"
        st.markdown(f"""
<div class="nv-card" style="height:100%;">
  <div class="nv-card-title">📈 Agent Confidence</div>
  <div style="font-size:2.4rem;font-weight:800;color:{conf_colour};margin-top:0.4rem;">
    {conf_pct}%
  </div>
  <div class="conf-bar-wrap">
    <div class="conf-bar-fill" style="width:{conf_pct}%;background:{'linear-gradient(90deg,#3fb950,#2ea043)' if conf_pct>=75 else 'linear-gradient(90deg,#d29922,#9e6a03)' if conf_pct>=50 else 'linear-gradient(90deg,#f85149,#da3633)'};"></div>
  </div>
</div>
""", unsafe_allow_html=True)

    # ── Tabs: Output / Critique / Log / Ethics ─────────────────────────────────────
    st.markdown("### 📄 Output")
    tab_out, tab_crit, tab_log, tab_eth = st.tabs([
        "✅ Refined Output", "🔎 Self-Critique", "🔄 Activity Log", "⚖️ Ethics & Compliance"
    ])

    with tab_out:
        mode_used = result.strategy_used
        is_diagram = st.session_state.diagram_mode
        st.markdown(
            f"<div style='font-size:0.75rem;color:#6e7b8b;margin-bottom:0.5rem;'>"
            f"{'📊 Diagram Mode' if is_diagram else '📝 Text Mode'} &nbsp;·&nbsp; "
            f"Strategy: <code>{mode_used}</code></div>",
            unsafe_allow_html=True
        )
        st.markdown(
            "<div style='background:#0d1117;border-left:3px solid #d2992255;"
            "border-radius:0 6px 6px 0;padding:0.45rem 0.9rem;margin-bottom:0.9rem;"
            "font-size:0.75rem;color:#9e7a2e;'>"
            "⚠️ <b>AI-generated content.</b> Review carefully before use — "
            "mathematical expressions, diagram values, and technical details may contain errors."
            "</div>",
            unsafe_allow_html=True
        )

        def _render_with_mermaid(md_text: str):
            pattern = re.compile(r'```mermaid\s*\n(.*?)\n```', re.DOTALL | re.IGNORECASE)
            parts = pattern.split(md_text)
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    if part.strip():
                        st.markdown(part)
                else:
                    safe = part.strip()
                    mermaid_html = f"""<!DOCTYPE html>
<html><head>
<script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
<style>
  body {{ margin:0; padding:0; background:#1e2130; }}
  .mermaid {{ background:#1e2130; padding:1rem; border-radius:10px; }}
  svg {{ max-width:100%; }}
  pre.fallback {{ background:#0d1117; color:#8b949e; font-size:0.75rem;
                  padding:1rem; border-radius:8px; overflow-x:auto;
                  border:1px solid #2d3250; margin:0; }}
</style>
</head><body>
<div id="graph" class="mermaid"></div>
<pre id="fallback" class="fallback" style="display:none"></pre>
<script>
const code = {repr(safe)};
document.getElementById('graph').textContent = code;
document.getElementById('fallback').textContent = code;
mermaid.initialize({{
  startOnLoad: false, theme: 'dark', securityLevel: 'loose',
  themeVariables: {{
    primaryColor: '#667eea', primaryTextColor: '#c9d1d9',
    primaryBorderColor: '#764ba2', lineColor: '#8b949e',
    background: '#1e2130', mainBkg: '#0d1117',
    nodeBorder: '#667eea', clusterBkg: '#161b27'
  }}
}});
mermaid.run().catch(function() {{
  document.getElementById('graph').style.display = 'none';
  document.getElementById('fallback').style.display = 'block';
}});
</script>
</body></html>"""
                    components.html(mermaid_html, height=380, scrolling=True)

        _render_with_mermaid(result.refined_markdown)

    with tab_crit:
        refined_used = ("looks complete and accurate" not in result.critique.lower())
        st.markdown('<div class="nv-card-title">Critique Result</div>', unsafe_allow_html=True)
        st.markdown(result.critique)
        note = "⚙️ Refinement pass was applied based on this critique." if refined_used else "✅ No refinement needed — output passed review."
        st.markdown(f'<div style="margin-top:0.4rem;font-size:0.78rem;color:#6e7b8b;">{note}</div>', unsafe_allow_html=True)

    with tab_log:
        log_html = "<br>".join(result.agent_log)
        st.markdown(f'<div class="agent-log">{log_html}</div>', unsafe_allow_html=True)

    with tab_eth:
        sr = result.safety_result or {}
        is_safe = sr.get("is_safe", True)
        concern = sr.get("concern_type", "none")
        explanation = sr.get("explanation", "Not screened.")
        safe_colour = "#3fb950" if is_safe else "#d29922"
        safe_label = "✓ No sensitive content detected" if is_safe else f"⚠️ Concern acknowledged: {concern.replace('_', ' ').title()}"
        ack_note = "" if is_safe else "<div style='font-size:0.78rem;color:#d29922;margin-top:0.5rem;'>⚠️ User provided explicit acknowledgment before processing proceeded (Human-in-the-Loop gate).</div>"
        st.markdown(f"""
<div class="nv-card">
  <div class="nv-card-title">🛡️ Safety Screening Result (Step 0 — Ethical Gate)</div>
  <div style="color:{safe_colour};font-weight:700;margin-bottom:0.4rem;">{safe_label}</div>
  <div style="font-size:0.83rem;color:#8b949e;">{explanation}</div>
  {ack_note}
</div>
<div class="nv-card">
  <div class="nv-card-title">⚖️ ACM/IEEE Code of Ethics — Principles Applied This Run</div>
  <div style="font-size:0.83rem;color:#8b949e;line-height:1.95;">
    <b style="color:#c9d1d9;">§1.1 Contribute to society</b> — Converts inaccessible handwriting into usable digital text<br>
    <b style="color:#c9d1d9;">§1.6 Respect privacy</b> — No image or personal data retained after session<br>
    <b style="color:#c9d1d9;">§2.5 Thorough evaluation</b> — Critique &amp; refinement loop validates quality before delivery<br>
    <b style="color:#c9d1d9;">§2.9 Design for inclusion</b> — Consent gate + safety gate prevent uninformed or harmful use<br>
    <b style="color:#c9d1d9;">§3.1 Know your limitations</b> — Confidence score transparently communicated; system is semi-autonomous only<br>
    <span style="color:#d29922;">⚠️ Potential §2.5 violation (model bias)</span> — LLM may misread handwriting styles/languages. Mitigated by confidence scoring, HITL approval gate, and refinement loop.
  </div>
</div>
<div class="nv-card">
  <div class="nv-card-title">🔒 Data Handling Transparency</div>
  <div style="font-size:0.83rem;color:#8b949e;line-height:1.75;">
    • Image processed by Google Gemini 2.5 Flash via Vertex AI (GCP, region: us-central1)<br>
    • No image data persisted beyond the API call lifetime (GDPR-aligned design)<br>
    • Session memory cleared on browser refresh — no long-term personal data retention<br>
    • Every agent decision is logged and auditable via the Activity Log tab<br>
    • HITL gate is mandatory — no output can be downloaded without explicit user approval
  </div>
</div>
<div class="nv-card">
  <div class="nv-card-title">⚠️ Responsible AI — Known Limitations &amp; Bias Disclosure</div>
  <div style="font-size:0.83rem;color:#8b949e;line-height:1.75;">
    • This is a <b style="color:#c9d1d9;">semi-autonomous system</b> — human oversight is required at all times<br>
    • LLM accuracy varies across handwriting styles, languages, and image quality<br>
    • Mathematical symbol recognition depends on legibility — always verify LaTeX output<br>
    • Confidence scores are model estimates, not guaranteed accuracy metrics<br>
    • Do not use this system as sole source of truth for academic or professional documents
  </div>
</div>
""", unsafe_allow_html=True)

    # ── Human-in-the-Loop Gate ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("""
<div class="hitl-banner">
  <div class="hitl-title">🟢 Human-in-the-Loop Control</div>
  <div class="hitl-desc">
    Review the agent's output above. You control what happens next — approve to download,
    reject to discard, or re-run for another agent pass.
  </div>
</div>
""", unsafe_allow_html=True)

    col_a, col_r, col_rr = st.columns([1, 1, 1])
    with col_a:
        if st.button("✅ Accept Output", use_container_width=True):
            st.session_state.hitl_accepted = True
            st.rerun()
    with col_r:
        if st.button("❌ Reject Output", use_container_width=True):
            st.session_state.hitl_accepted = False
            st.session_state.agent_result = None
            st.rerun()
    with col_rr:
        if st.button("🔁 Re-run", use_container_width=True):
            st.session_state.agent_result = None
            st.session_state.hitl_accepted = None
            st.rerun()

    # ── Download Section (only if accepted) ───────────────────────────────────
    if st.session_state.hitl_accepted is True:
        st.success("✅ Output accepted! Downloads are now available.")
        st.markdown("### 💾 Download")

        final_md = result.refined_markdown

        def _make_docx(md_text: str) -> bytes:
            doc = Document()
            t = doc.add_paragraph("Converted Notes")
            t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            t.runs[0].font.size = Pt(18)
            t.runs[0].font.bold = True
            doc.add_paragraph()
            for line in md_text.split("\n"):
                if line.strip():
                    if line.startswith("# "):
                        p = doc.add_paragraph(line[2:])
                        p.runs[0].font.size = Pt(16)
                        p.runs[0].font.bold = True
                    elif line.startswith("## "):
                        p = doc.add_paragraph(line[3:])
                        p.runs[0].font.size = Pt(14)
                        p.runs[0].font.bold = True
                    elif line.startswith("### "):
                        p = doc.add_paragraph(line[4:])
                        p.runs[0].font.size = Pt(12)
                        p.runs[0].font.bold = True
                    elif line.strip().startswith(("-", "*")):
                        doc.add_paragraph(line.strip()[1:].strip(), style="List Bullet")
                    else:
                        doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        col_d1, col_d2, _ = st.columns([1, 1, 2])
        with col_d1:
            st.download_button(
                "📥 DOCX File",
                data=_make_docx(final_md),
                file_name="notevision_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_d2:
            st.download_button(
                "📥 Markdown File",
                data=final_md,
                file_name="notevision_output.md",
                mime="text/markdown",
                use_container_width=True,
            )

    elif st.session_state.hitl_accepted is False:
        st.warning("Output rejected. Upload a new image or adjust the settings and re-run.")

# ─────────────────────────────────────────────────────────────────────────────
# Footer
# ─────────────────────────────────────────────────────────────────────────────

st.markdown("---")
st.markdown("""
<div style="text-align:center;padding:1.5rem 0;">
  <p style="color:#4a5568;font-size:0.85rem;margin:0;">
    NoteVision &nbsp;·&nbsp; Intelligent Handwriting Converter
  </p>
  <p style="color:#3d4559;font-size:0.75rem;margin-top:6px;">
    Powered by Google Gemini 2.5 Flash via Vertex AI &nbsp;|&nbsp; Built with Streamlit
  </p>
</div>
""", unsafe_allow_html=True)
